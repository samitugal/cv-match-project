import logging
import os
from typing import Any, Dict, List, Optional

import pytesseract
from docx import Document
from dotenv import load_dotenv
from pdf2image import convert_from_path
from PyPDF2 import PdfReader

from src.helpers.string_ops import strip_accents, mask_likely_personal_info

load_dotenv()

logger = logging.getLogger(__name__)


class DocxReader:
    def __init__(self, file_path: str):
        self.file_path = file_path
        self.document = None
        self._load_document()

    def _load_document(self):
        try:
            self.document = Document(self.file_path)
            logger.info(f"Word document loaded successfully: {self.file_path}")
        except Exception as e:
            logger.error(f"Failed to load Word document: {e}")
            raise

    def extract_text(self) -> str:
        if not self.document:
            return ""

        text_content = []
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)

        return "\n".join(text_content)

    def extract_paragraphs(self) -> List[str]:
        if not self.document:
            return []

        paragraphs = []
        for paragraph in self.document.paragraphs:
            if paragraph.text.strip():
                paragraphs.append(paragraph.text)

        return paragraphs

    def extract_tables(self) -> List[List[List[str]]]:
        if not self.document:
            return []

        tables_data = []
        for table in self.document.tables:
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            tables_data.append(table_data)

        return tables_data

    def extract_all_content(self) -> Dict[str, Any]:
        return {
            "text": self.extract_text(),
            "paragraphs": self.extract_paragraphs(),
            "tables": self.extract_tables(),
        }

    def get_document_info(self) -> Dict[str, Any]:
        if not self.document:
            return {}

        core_properties = self.document.core_properties
        return {
            "title": core_properties.title,
            "author": core_properties.author,
            "subject": core_properties.subject,
            "created": core_properties.created,
            "modified": core_properties.modified,
            "paragraphs_count": len(self.document.paragraphs),
            "tables_count": len(self.document.tables),
        }


def read_pdf_with_ocr(file_path: str) -> str:
    try:
        pytesseract.pytesseract.tesseract_cmd = os.getenv("TESERRACT_EXE_PATH")
        images = convert_from_path(
            file_path, dpi=400, poppler_path=os.getenv("POPPLER_BIN_PATH")
        )

        full_text = ""
        for image in images:
            text = pytesseract.image_to_string(image, lang="eng")
            normalized_text = strip_accents(text)
            full_text += normalized_text + "\n"

        return full_text.strip()
    except Exception as e:
        logger.error(f"OCR Error: {str(e)}")
        return ""


def read_pdf(file_path: str) -> Optional[str]:
    try:
        reader = PdfReader(file_path)
        full_text = ""
        for page in reader.pages:
            extracted = page.extract_text()
            if extracted:
                normalized = strip_accents(extracted)
                full_text += normalized + "\n"
        return full_text.strip()
    except Exception as e:
        logger.error(f"Error reading PDF: {str(e)}")
        return None


def get_page_count(file_path: str) -> Optional[int]:
    try:
        reader = PdfReader(file_path)
        return len(reader.pages)
    except Exception as e:
        logger.error(f"Error getting page count: {str(e)}")
        return None


def read_docx(file_path: str) -> Dict[str, Any]:
    reader = DocxReader(file_path)
    return reader.extract_all_content()


def read_document(file_path: str) -> Dict[str, Any]:
    if file_path.endswith(".pdf"):
        return read_pdf_with_ocr(file_path)
    elif file_path.endswith(".docx"):
        return read_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {file_path}")
